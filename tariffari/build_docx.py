# -*- coding: utf-8 -*-
"""Versione Word editabile del Tariffario prosecuzione sfratto — Studio Ansalone."""
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERDE   = RGBColor(0x1E, 0x4A, 0x38)
INK     = RGBColor(0x22, 0x22, 0x22)
GRIGIO  = RGBColor(0x6B, 0x6B, 0x6B)
BIANCO  = RGBColor(0xFF, 0xFF, 0xFF)
NOTE    = RGBColor(0x44, 0x44, 0x44)
VERDE_HEX  = "1E4A38"
BEIGE_HEX  = "F4F2EC"
BORDO_HEX  = "D8D3C4"

DEST = r"C:\Users\danil\Dropbox\claude_projects\sfratto-sicuro\tariffari"
IMG  = DEST + r"\monogramma-da-verde.png"
OUT  = DEST + r"\Tariffario-Sfratto-Prosecuzione.docx"

doc = Document()

# ---------- pagina ----------
sec = doc.sections[0]
sec.page_width  = Mm(210)
sec.page_height = Mm(297)
sec.top_margin    = Cm(1.5)
sec.bottom_margin = Cm(1.4)
sec.left_margin   = Cm(2.0)
sec.right_margin  = Cm(2.0)

# ---------- stile Normale ----------
normal = doc.styles["Normal"]
normal.font.name = "Garamond"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.04
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), "Garamond")
rfonts.set(qn("w:hAnsi"), "Garamond")
rfonts.set(qn("w:cs"), "Garamond")

# ---------- helpers ----------
def shade(cell, hexc):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(shd)

def borders(table, hexc, sz=8, edges=("top","left","bottom","right","insideH","insideV")):
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in edges:
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), hexc)
        b.append(el)
    tblPr.append(b)

def no_borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{e}"); el.set(qn("w:val"), "nil")
        b.append(el)
    tblPr.append(b)

def cell_pad(cell, top=60, bottom=60, left=120, right=120):  # twips
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)

def set_w(cell, cm):
    cell.width = Cm(cm)

def para_bottom_rule(p, hexc, sz=4):
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), hexc)
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)

def spacing(run, pts):
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(int(pts * 20)))
    run._r.get_or_add_rPr().append(sp)

def R(p, text, size=10.5, color=INK, bold=False, italic=False, sc=False, track=None):
    r = p.add_run(text)
    r.font.name = "Garamond"; r.font.size = Pt(size)
    r.font.color.rgb = color; r.bold = bold; r.italic = italic
    if sc: r.font.small_caps = True
    if track: spacing(r, track)
    rp = r._r.get_or_add_rPr(); rf = rp.get_or_add_rFonts()
    rf.set(qn("w:ascii"), "Garamond"); rf.set(qn("w:hAnsi"), "Garamond"); rf.set(qn("w:cs"), "Garamond")
    return r

def bullet(cell, text, size=10.5, color=INK, after=2):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    R(p, text, size=size, color=color)
    return p

# ---------- CARTA INTESTATA ----------
lt = doc.add_table(rows=1, cols=2)
no_borders(lt); lt.autofit = False
lc, rc = lt.rows[0].cells
set_w(lc, 3.5); set_w(rc, 13.5)
lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
lp = lc.paragraphs[0]; lp.add_run().add_picture(IMG, height=Mm(15))
recap = [
    ("Avv. Danilo Ansalone", True),
    ("Via F. P. Volpe n. 2 — 84122 Salerno", False),
    ("Tel +39 089 0978260 — consulenza@daniloansalone.it", False),
    ("PEC danilo.ansalone@pec.giuffre.it — www.daniloansalone.it", False),
    ("Ordine degli Avvocati di Salerno", False),
]
rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.paragraph_format.space_after = Pt(0); rp.paragraph_format.line_spacing = 1.0
for i, (txt, sc) in enumerate(recap):
    if i > 0: rp.add_run().add_break()
    R(rp, txt, size=9, sc=sc, track=(1.0 if sc else None))

rule = doc.add_paragraph(); rule.paragraph_format.space_before = Pt(2); rule.paragraph_format.space_after = Pt(2)
para_bottom_rule(rule, "222222", sz=4)

# ---------- TITOLO ----------
ov = doc.add_paragraph(); ov.alignment = WD_ALIGN_PARAGRAPH.CENTER
ov.paragraph_format.space_before = Pt(6); ov.paragraph_format.space_after = Pt(2)
R(ov, "TARIFFARIO DEI COMPENSI PROFESSIONALI", size=9, color=VERDE, bold=True, track=2.0)

tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
tt.paragraph_format.space_after = Pt(2)
R(tt, "Prosecuzione dello sfratto dopo l’opposizione", size=18, color=VERDE, bold=True)

st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.paragraph_format.space_after = Pt(6)
R(st, "Fase di merito successiva al mutamento del rito (art. 667 c.p.c.)", size=10.5, color=GRIGIO, italic=True)

# ---------- INTRO ----------
intro = doc.add_paragraph(); intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro.paragraph_format.space_after = Pt(8)
R(intro, "Il compenso di ")
R(intro, "€ 1.300,00", bold=True)
R(intro, " (IVA e CPA comprese) riguarda la ")
R(intro, "fase iniziale", italic=True)
R(intro, " della procedura di sfratto — intimazione, notifica e udienza di convalida. "
         "In caso di opposizione dell’inquilino e prosecuzione del giudizio dopo il mutamento del rito, "
         "la causa entra nella fase di merito e si applicano, ")
R(intro, "in aggiunta al compenso della fase iniziale", bold=True)
R(intro, ", le tariffe indicate di seguito.")

# ---------- FASE INIZIALE ----------
fi = doc.add_table(rows=1, cols=2); borders(fi, BORDO_HEX, sz=4); fi.autofit = False
c0, c1 = fi.rows[0].cells
set_w(c0, 12.0); set_w(c1, 5.0)
for c in (c0, c1):
    shade(c, BEIGE_HEX); cell_pad(c); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = c0.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
R(p, "Fase iniziale", bold=True); R(p, " — intimazione, notifica, udienza di convalida")
p = c1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after = Pt(0)
R(p, "€ 1.300,00", size=14, color=VERDE, bold=True)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- FASCIA DI PREZZO ----------
def fascia(titolo, prezzo, condizione, occhiello, voci):
    t = doc.add_table(rows=2, cols=2); borders(t, VERDE_HEX, sz=10); t.autofit = False
    h0, h1 = t.rows[0].cells
    set_w(h0, 12.0); set_w(h1, 5.0)
    for c in (h0, h1):
        shade(c, VERDE_HEX); cell_pad(c, top=80, bottom=80); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = h0.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    R(p, titolo, size=13, color=BIANCO, bold=True)
    p = h1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after = Pt(0)
    R(p, prezzo, size=16, color=BIANCO, bold=True)
    body = t.rows[1].cells[0].merge(t.rows[1].cells[1])
    cell_pad(body, top=120, bottom=140, left=140, right=140)
    bp = body.paragraphs[0]; bp.paragraph_format.space_after = Pt(4)
    R(bp, condizione, size=10, color=GRIGIO, italic=True)
    op = body.add_paragraph(); op.paragraph_format.space_after = Pt(2); op.paragraph_format.space_before = Pt(2)
    R(op, occhiello, size=9.5, color=VERDE, bold=True, sc=True, track=0.5)
    for v in voci:
        bullet(body, v)
    return t

fascia(
    "Tariffa base", "€ 2.500,00",
    "Quando la causa non richiede prove testimoniali, consulenza tecnica o altra attività istruttoria.",
    "Comprende",
    ["studio dell’opposizione;",
     "redazione e deposito delle memorie;",
     "procedimento di mediazione obbligatoria;",
     "partecipazione alle udienze;",
     "attività difensiva fino alla sentenza."],
)
doc.add_paragraph().paragraph_format.space_after = Pt(3)
fascia(
    "Tariffa completa", "€ 3.500,00",
    "Quando il giudizio richiede attività istruttoria, in aggiunta alla trattazione.",
    "In particolare quando ricorrono",
    ["prove testimoniali;",
     "interrogatorio formale delle parti;",
     "consulenza tecnica d’ufficio (CTU);",
     "più udienze istruttorie;",
     "attività difensiva particolarmente complessa."],
)

# ---------- ESCLUSIONI ----------
exh = doc.add_paragraph(); exh.paragraph_format.space_before = Pt(8); exh.paragraph_format.space_after = Pt(2)
R(exh, "RESTANO ESCLUSE (A CARICO DEL CLIENTE)", size=11, color=VERDE, bold=True, sc=True, track=0.5)
et = doc.add_table(rows=1, cols=2); no_borders(et); et.autofit = False
ecL, ecR = et.rows[0].cells
set_w(ecL, 8.5); set_w(ecR, 8.5)
ecL.paragraphs[0]._p.getparent().remove(ecL.paragraphs[0]._p)  # rimuove il paragrafo vuoto iniziale
ecR.paragraphs[0]._p.getparent().remove(ecR.paragraphs[0]._p)
for v in ["spese dell’organismo di mediazione;",
          "contributo unificato e marca da bollo;",
          "spese di notifica;",
          "compensi di consulenti tecnici (CTU / CTP);"]:
    bullet(ecL, v)
for v in ["spese vive e di trasferta;",
          "attività esecutiva successiva alla sentenza (preavviso di rilascio, ufficiale giudiziario);",
          "giudizio di appello."]:
    bullet(ecR, v)

sep = doc.add_paragraph(); sep.paragraph_format.space_before = Pt(6); sep.paragraph_format.space_after = Pt(4)
para_bottom_rule(sep, BORDO_HEX, sz=4)

# ---------- NOTE ----------
note_items = [
    [("I compensi indicati si riferiscono al ", False), ("solo primo grado", True), (" di giudizio.", False)],
    [("Il compenso è dovuto ", False), ("indipendentemente dall’esito", True),
     (" della causa: l’avvocato assume un’obbligazione di mezzi, non di risultato.", False)],
    [("Gli importi sono ", False), ("comprensivi di IVA e contributo Cassa Forense (CPA)", True),
     (" e sono al netto delle spese sopra escluse.", False)],
    [("Il presente tariffario vale come preventivo di massima ai sensi dell’art. 13, co. 5, L. 31 dicembre 2012 n. 247; "
      "per ciascun incarico verrà rilasciato preventivo scritto personalizzato.", False)],
]
for parts in note_items:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.line_spacing = 1.0
    R(p, "▪  ", size=9.5, color=VERDE)
    for txt, b in parts:
        R(p, txt, size=9.5, color=NOTE, bold=b)

# ---------- FIRMA ----------
ft = doc.add_table(rows=1, cols=2); no_borders(ft); ft.autofit = False
fl, fr = ft.rows[0].cells
set_w(fl, 9.0); set_w(fr, 8.0)
fl.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
fr.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
p = fl.paragraphs[0]; p.paragraph_format.space_before = Pt(6)
R(p, "Listino valido dal 13 luglio 2026 — Salerno.", size=9, color=GRIGIO, italic=True)
p = fr.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_before = Pt(6)
R(p, "Studio Legale Ansalone", size=10); p.add_run().add_break()
R(p, "Avv. Danilo Ansalone", size=10, bold=True)

# ---------- FOOTER ----------
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_bottom_rule(fp, "CCCCCC", sz=4)  # visivo: linea sopra? no, mettiamo top
# usiamo bordo superiore per il footer
pPr = fp._p.get_or_add_pPr()
# rimuovi eventuale pBdr appena aggiunto e aggiungi bordo TOP
for pbdr in pPr.findall(qn("w:pBdr")):
    pPr.remove(pbdr)
pbdr = OxmlElement("w:pBdr"); top = OxmlElement("w:top")
top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "4"); top.set(qn("w:space"), "2"); top.set(qn("w:color"), "CCCCCC")
pbdr.append(top); pPr.append(pbdr)
R(fp, "Studio Legale Ansalone — C.F. NSLDNL75S04H703U · P. IVA 04380590655", size=8, color=GRIGIO)

doc.save(OUT)
print("OK ->", OUT)
